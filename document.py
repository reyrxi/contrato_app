import os

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def _replace_in_paragraph(paragraph, mapping):
    for run in paragraph.runs:
        for placeholder, value in mapping.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    for placeholder, value in mapping.items():
        while True:
            texts = [r.text for r in paragraph.runs]
            combined = "".join(texts)
            if placeholder not in combined:
                break
            if any(placeholder in t for t in texts):
                break
            start = combined.index(placeholder)
            end = start + len(placeholder)
            pos = 0
            start_run = end_run = -1
            for i, t in enumerate(texts):
                if start_run == -1 and pos + len(t) > start:
                    start_run = i
                if pos + len(t) >= end:
                    end_run = i
                    break
                pos += len(t)
            if start_run == -1 or end_run == -1 or start_run == end_run:
                break
            merged = "".join(r.text for r in paragraph.runs[start_run:end_run + 1])
            paragraph.runs[start_run].text = merged.replace(placeholder, value)
            for run in paragraph.runs[start_run + 1:end_run + 1]:
                run.text = ""
            break


def _replace_in_element(element, mapping):
    for para_el in element.iter(qn("w:p")):
        _replace_in_paragraph(Paragraph(para_el, None), mapping)


def replace_all(doc, mapping):
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)

    for text_box in doc.element.body.iter(qn("w:txbxContent")):
        _replace_in_element(text_box, mapping)

    for section in doc.sections:
        header_footers = [
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ]
        for hf in header_footers:
            if not hf:
                continue
            for paragraph in hf.paragraphs:
                _replace_in_paragraph(paragraph, mapping)
            for text_box in hf._element.iter(qn("w:txbxContent")):
                _replace_in_element(text_box, mapping)


def convert_to_pdf(docx_path, pdf_path):
    try:
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        return True, ""
    except ImportError:
        pass

    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return True, ""
    except ImportError:
        pass

    return False, "Instale comtypes ou docx2pdf para gerar PDF."
